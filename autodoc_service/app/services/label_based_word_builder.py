"""
Label-Based Word Document Builder

This module implements a robust label-based mapping system that finds label text
in Word templates and fills adjacent cells with data. This approach is much more
resilient to template structure changes than cell index-based mapping.

Key Benefits:
- Template structure independent (works with .docx, .docm, different layouts)
- Finds labels by text content and fills the next cell in the same row
- Handles multiline content properly
- Robust against cell merging and table restructuring
"""
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

from docx import Document
from docx.table import Table, _Row, _Cell
from docx.shared import Pt

from ..models import ChangeRequest
from .paths import verify_template_exists, get_documents_dir
from .filename import generate_word_filename, unique_path
from .word_payload import build_word_payload
from .font_styler import ensure_malgun_gothic_document
from ..logging_config import get_logger

# 모듈 로거 설정
logger = get_logger(__name__)


def normalize_label(text: str) -> str:
    """
    Normalize label text for consistent matching
    
    Removes whitespace, newlines, colons, and other formatting characters
    to enable reliable label matching regardless of template formatting.
    """
    if not text:
        return ""
    
    # Remove extra whitespace, newlines, colons, and common punctuation
    normalized = re.sub(r'[\s\n\r:：]+', '', text.strip())
    
    # Convert to lowercase for case-insensitive matching
    normalized = normalized.lower()
    
    return normalized


def create_label_mapping() -> Dict[str, str]:
    """
    Create mapping from normalized template labels to JSON data keys
    
    This mapping defines which data field corresponds to each label
    found in the Word template.
    """
    return {
        # Basic Information
        '제목': '제목',
        '변경관리번호': '변경관리번호', 
        '작업일시': '작업일시',
        '배포일시': '배포일시',
        
        # Customer and Request Info
        '고객사명': '고객사',
        '요청부서': '요청부서',
        '요청자': '요청자',
        '요청번호(sor)': '문서번호',
        '요청번호': '문서번호',  # Alternative label
        
        # System Information
        '대상시스템': '요청시스템',
        '시스템': '요청시스템',  # Alternative label
        
        # Personnel
        '작업자/배포자': '작업자-배포자',
        '작업자배포자': '작업자-배포자',  # No slash variant
        '처리자': '처리자_약칭',
        
        # Content
        '목적/개선내용': '목적-개선내용',
        '목적개선내용': '목적-개선내용',  # No slash variant
        '영향도대상자': '영향도_대상자',
        '영향도대상자': '영향도_대상자',  # Alternative spelling
        
        # Additional fields that might be in template
        '완료희망일': '완료희망일',
        '개발등급': '개발등급',
        '작업예상일자': '작업예상일자',
        
        # Date fields
        '기안': '작성일_mmdd',
        '작성일': '작성일_mmdd',
    }


def find_label_cell_in_row(row: _Row, normalized_label: str) -> Optional[_Cell]:
    """
    Find a cell in the row that contains the specified normalized label
    
    Returns the cell containing the label, or None if not found.
    """
    for cell in row.cells:
        cell_text = normalize_label(cell.text)
        if normalized_label in cell_text or cell_text in normalized_label:
            return cell
    return None


def get_data_cell_for_label(row: _Row, label_cell: _Cell, label_text: str) -> Optional[_Cell]:
    """
    Get the appropriate data cell for a label based on template patterns
    
    This function uses template-specific logic to identify where data should go
    for each label, handling merged cells and complex table structures.
    """
    try:
        cells = list(row.cells)
        label_normalized = normalize_label(label_text)
        
        # Debug output (can be removed in production)
        # print(f"      🔧 get_data_cell_for_label: '{label_text}' → '{label_normalized}' (total cells: {len(cells)})")
        
        # Template-specific data cell patterns based on the actual structure
        if label_normalized == "제목":
            # 제목: Cell 1 (label) → Cell 2 (data)
            result = cells[1] if len(cells) > 1 else None
            # print(f"      📋 제목 logic → Cell 2: {result is not None}")
            return result
            
        elif label_normalized == "변경관리번호":
            # 변경관리번호: Cell 11/12/13 (label) → Cell 14 (data)
            return cells[13] if len(cells) > 13 else None
            
        elif label_normalized == "작업일시":
            # 작업일시: Cell 1 (label) → Cell 2 (data) 
            return cells[1] if len(cells) > 1 else None
            
        elif label_normalized == "배포일시":
            # 배포일시: Cell 8/9/10 (label) → Cell 11 (data)
            return cells[10] if len(cells) > 10 else None
            
        elif label_normalized == "고객사명":
            # 고객사명: Cell 1 (label) → Cell 2 (data)
            return cells[1] if len(cells) > 1 else None
            
        elif label_normalized == "요청부서":
            # 요청부서: Cell 7 (label) → Cell 8 (data)
            return cells[7] if len(cells) > 7 else None
            
        elif label_normalized == "요청자":
            # 요청자: Cell 11/12 (label) → Cell 13 (data) 
            return cells[12] if len(cells) > 12 else None
            
        elif label_normalized == "요청번호(sor)":
            # 요청번호(SOR): Cell 11/12 (label) → Cell 13 (data)
            return cells[12] if len(cells) > 12 else None
            
        elif label_normalized == "대상시스템":
            # 대상 시스템: Cell 1 (label) → Cell 2 (data)
            return cells[1] if len(cells) > 1 else None
            
        elif label_normalized == "작업자/배포자":
            # 작업자/배포자: Cell 8/9/10 (label) → Cell 11 (data)
            return cells[10] if len(cells) > 10 else None
            
        elif label_normalized == "목적/개선내용":
            # 목적/개선내용: Table 3, Row 11, Cell 1 (label) → Cell 2 (data)
            return cells[1] if len(cells) > 1 else None
            
        # 기안 필드는 handle_table2_special에서 별도 처리됨
        
        else:
            # Default: try to find label_cell index and use next cell
            try:
                label_index = cells.index(label_cell)
                if label_index + 1 < len(cells):
                    return cells[label_index + 1]
            except ValueError:
                pass
                
    except (IndexError):
        pass
    return None


def set_cell_content(cell: _Cell, content: str):
    """
    Set cell content with proper multiline handling
    
    Note: Font styling is applied later via ensure_malgun_gothic_document()
    to ensure document-wide consistency.
    """
    if not content:
        content = ""
    
    # Clear existing content
    cell.text = ""
    
    # Handle multiline content
    if '\n' in content:
        lines = content.split('\n')
        paragraph = cell.paragraphs[0]
        
        for i, line in enumerate(lines):
            if i > 0:
                # Add line break for subsequent lines
                run = paragraph.add_run()
                run.add_break()
            paragraph.add_run(line)
    else:
        # Single line content
        cell.text = str(content)


def handle_table2_special(table: Table, data: Dict[str, Any], filled_data_keys: set) -> int:
    """
    Table 2의 특별한 구조를 처리하는 함수
    
    Table 2 구조:
    - Row 1: 기안 | 검토 | 검토 | 검토 | ERD변경 | IF관리 | 승인 | 의견
    - Row 2: 홍길동 | 김태우 | | | | | 이동호 | 
    - Row 3: ~ 3/25(3/28) | / | / | / | / | / | / | 
    
    기안 날짜는 Row 3, Cell 1에 들어가야 함
    """
    filled_count = 0
    
    if len(table.rows) >= 3:  # Table 2에 최소 3개 행이 있어야 함
        # 기안 날짜 처리 (Row 3, Cell 1)
        date_value = data.get('작성일_mmdd', '')
        if date_value and '작성일_mmdd' not in filled_data_keys:
            try:
                date_cell = table.rows[2].cells[0]  # Row 3, Cell 1 (0-indexed)
                set_cell_content(date_cell, str(date_value))
                filled_count += 1
                filled_data_keys.add('작성일_mmdd')
                print(f"✅ Table 2, Row 3, Cell 1: '기안 날짜' → '작성일_mmdd': {date_value}...")
            except Exception as e:
                print(f"⚠️ Failed to fill Table 2 기안 날짜: {e}")
    
    return filled_count


def fill_template_by_labels(doc: Document, data: Dict[str, Any]) -> int:
    """
    Fill Word template using label-based mapping
    
    Searches all tables for label cells and fills adjacent cells with data.
    Uses improved matching to prevent duplicate fills and ensure accuracy.
    
    Returns:
        int: Number of fields successfully mapped
    """
    label_mapping = create_label_mapping()
    filled_count = 0
    filled_data_keys = set()  # Track filled data keys to prevent duplicates
    
    print(f"🔍 Available data keys: {list(data.keys())}")
    print(f"🗝️ Mapping {len(label_mapping)} labels...")
    
    # Process all tables in the document
    for table_idx, table in enumerate(doc.tables):
        print(f"\n📋 Processing Table {table_idx + 1}...")
        
        # Special handling for Table 2 (기안 날짜 문제)
        if table_idx == 1:  # Table 2 (0-indexed)
            filled_count += handle_table2_special(table, data, filled_data_keys)
            continue
        
        for row_idx, row in enumerate(table.rows):
            
            # Check each label in our mapping
            for normalized_label, data_key in label_mapping.items():
                
                # Skip if we already filled this data key
                if data_key in filled_data_keys:
                    continue
                
                # Check if we have data for this key
                data_value = data.get(data_key, "")
                if not data_value:  # Skip empty values
                    continue
                
                # Find the label cell in this row that exactly matches
                label_cell = None
                for cell in row.cells:
                    cell_normalized = normalize_label(cell.text)
                    
                    # More precise matching to avoid false positives
                    if cell_normalized == normalized_label:
                        # Exact match is best
                        if len(cell.text.strip()) > 1:  # Must have meaningful content
                            label_cell = cell
                            break
                    elif (len(normalized_label) > 4 and normalized_label in cell_normalized and
                          len(cell_normalized) <= len(normalized_label) + 2):
                        # Close match but avoid short/generic matches
                        if len(cell.text.strip()) > 2:
                            label_cell = cell
                            break
                
                if not label_cell:
                    continue
                
                # Get the appropriate data cell for this label
                data_cell = get_data_cell_for_label(row, label_cell, label_cell.text)
                if not data_cell or data_cell == label_cell:  # Ensure different cell
                    continue
                
                # Set the content
                try:
                    set_cell_content(data_cell, str(data_value))
                    filled_count += 1
                    filled_data_keys.add(data_key)  # Mark data key as filled
                    
                    print(f"✅ Table {table_idx + 1}, Row {row_idx + 1}: '{label_cell.text.strip()}' → '{data_key}': {str(data_value)[:50]}...")
                    
                except Exception as e:
                    print(f"⚠️ Failed to fill '{normalized_label}': {e}")
    
    print(f"\n📊 Summary: Filled {filled_count} unique fields")
    print(f"📋 Filled data keys: {sorted(filled_data_keys)}")
    return filled_count


def build_change_request_doc_label_based(
    data: ChangeRequest, 
    out_dir: Optional[Path] = None, 
    raw_data: Optional[dict] = None
) -> Path:
    """
    Build change request Word document using label-based mapping
    
    This function replaces the cell index-based approach with a robust
    label-based system that finds labels in the template and fills
    adjacent cells with data.
    
    Args:
        data: ChangeRequest model instance
        out_dir: Output directory (defaults to documents directory)
        raw_data: Raw parsed data for enhanced field mapping
        
    Returns:
        Path: Generated document file path
    """
    start_time = time.time()
    
    logger.info(f"Word 문서 생성 시작: change_id={data.change_id}, title={data.title}, has_raw_data={raw_data is not None}")
    
    # Validate required data
    if not data.change_id:
        logger.error("Word 문서 생성 실패: change_id 누락")
        raise ValueError("change_id는 필수입니다")
    if not data.title:
        logger.error("Word 문서 생성 실패: title 누락")
        raise ValueError("title은 필수입니다")
    
    # Load template
    try:
        template_path = verify_template_exists("template.docx")
        logger.info(f"템플릿 로드 성공: {template_path}")
        doc = Document(str(template_path))
    except Exception as e:
        logger.exception(f"템플릿 로드 실패: {str(e)}")
        raise
    
    # Set output directory
    if out_dir is None:
        out_dir = get_documents_dir()
    
    # Prepare data for mapping
    if raw_data:
        # Use enhanced payload with missing field derivation
        logger.info("향상된 페이로드 사용으로 누락 필드 자동 보완")
        word_payload = build_word_payload(raw_data)
        logger.info(f"향상된 페이로드 생성 완료: field_count={len(word_payload)}")
    else:
        # Convert ChangeRequest to dict for compatibility
        logger.info("기본 ChangeRequest 데이터로 페이로드 생성")
        word_payload = {
            'change_id': data.change_id,
            'title': data.title,
            'system': data.system,
            'requester': data.requester,
            'writer_short': data.writer_short,
            'work_datetime': data.work_datetime or "",
            'deploy_datetime': data.deploy_datetime or "",
            'customer': data.customer or "",
            'request_dept': data.request_dept or "",
            'doc_no': data.doc_no or "",
            'worker_deployer': data.worker_deployer or "",
            'impact_targets': data.impact_targets or "",
            'created_date': data.created_date or datetime.now().strftime("%m/%d"),
        }
        
        # Add purpose content
        if data.details and data.details.summary:
            word_payload['목적-개선내용'] = data.details.summary
        elif data.details and data.details.plan:
            word_payload['목적-개선내용'] = data.details.plan
        
        logger.info(f"기본 페이로드 생성 완료: field_count={len(word_payload)}")
    
    # Map field names to match template labels
    template_data = {
        '제목': word_payload.get('제목', data.title),
        '변경관리번호': word_payload.get('변경관리번호', data.change_id),
        '작업일시': word_payload.get('작업일시', ''),
        '배포일시': word_payload.get('배포일시', ''),
        '고객사': word_payload.get('고객사', word_payload.get('customer', '')),
        '요청부서': word_payload.get('요청부서', word_payload.get('request_dept', '')),
        '요청자': word_payload.get('요청자', data.requester),
        '문서번호': word_payload.get('문서번호', word_payload.get('doc_no', '')),
        '요청시스템': word_payload.get('요청시스템', data.system),
        '작업자-배포자': word_payload.get('작업자-배포자', word_payload.get('worker_deployer', '')),
        '목적-개선내용': word_payload.get('목적-개선내용', ''),
        '영향도_대상자': word_payload.get('영향도_대상자', word_payload.get('impact_targets', '')),
        '처리자_약칭': word_payload.get('처리자_약칭', data.writer_short),
        '작성일_mmdd': word_payload.get('작성일_mmdd', ''),
    }
    
    # Fill template using label-based mapping
    logger.info("라벨 기반 템플릿 채우기 시작")
    template_field_count = len(template_data)
    filled_count = fill_template_by_labels(doc, template_data)
    
    logger.info(f"템플릿 채우기 완료: filled_fields={filled_count}/{template_field_count}")
    
    # Apply 맑은 고딕 font to entire document (ensures consistency)
    logger.info("맑은 고딕 폰트 일관성 적용 시작")
    ensure_malgun_gothic_document(doc)
    logger.info("폰트 스타일링 완료")
    
    # Generate filename
    filename = generate_word_filename(
        change_id=data.change_id,
        title=data.title,
        writer_short=data.writer_short
    )
    logger.info(f"파일명 생성 완료: {filename}")
    
    # Create unique output path
    output_path = unique_path(out_dir, filename)
    logger.info(f"출력 경로 설정: {output_path}")
    
    # Save document
    try:
        doc.save(str(output_path))
        file_size = output_path.stat().st_size if output_path.exists() else 0
        processing_time = time.time() - start_time
        
        logger.info(f"Word 문서 생성 성공: filename={output_path.name}, size={file_size} bytes, processing_time={processing_time:.3f}s, filled_fields={filled_count}")
    except Exception as e:
        logger.exception(f"Word 문서 저장 실패: filename={filename}, error={str(e)}")
        raise
    
    return output_path