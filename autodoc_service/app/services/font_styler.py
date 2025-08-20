"""
Word 문서 폰트 스타일링 서비스

전체 Word 문서에 일관된 폰트(맑은 고딕)를 적용하는 기능을 제공합니다.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def apply_malgun_gothic_to_document(doc: Document) -> None:
    """
    Word 문서 전체에 맑은 고딕 폰트를 적용
    
    템플릿 원본 텍스트와 매핑된 데이터 모두에 일관된 폰트를 적용합니다.
    
    Args:
        doc: python-docx Document 객체
    """
    print("Applying 맑은 고딕 font to entire document...")
    
    # 1. 문서의 모든 단락(Paragraph)에 폰트 적용
    for paragraph in doc.paragraphs:
        apply_malgun_gothic_to_paragraph(paragraph)
    
    # 2. 모든 테이블의 셀에 폰트 적용
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_malgun_gothic_to_paragraph(paragraph)
    
    # 3. 헤더와 푸터에 폰트 적용
    for section in doc.sections:
        # 헤더 처리
        if section.header:
            for paragraph in section.header.paragraphs:
                apply_malgun_gothic_to_paragraph(paragraph)
        
        # 푸터 처리
        if section.footer:
            for paragraph in section.footer.paragraphs:
                apply_malgun_gothic_to_paragraph(paragraph)
    
    print("Document font styling completed")


def apply_malgun_gothic_to_paragraph(paragraph) -> None:
    """
    단락의 모든 Run에 맑은 고딕 폰트 적용
    
    Args:
        paragraph: python-docx Paragraph 객체
    """
    for run in paragraph.runs:
        if run.font:
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10)  # 기본 크기
            
            # 한글 폰트 설정 (중요: Word에서 한글 폰트 인식을 위함)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


def ensure_malgun_gothic_document(doc: Document) -> None:
    """
    Word 문서에 맑은 고딕 폰트가 일관되게 적용되도록 보장
    
    매핑된 데이터와 템플릿 텍스트 모두에 동일한 폰트를 적용합니다.
    
    Args:
        doc: python-docx Document 객체
    """
    apply_malgun_gothic_to_document(doc)


def apply_malgun_gothic_to_cell(cell, content: str = None) -> None:
    """
    특정 셀에 맑은 고딕 폰트 적용 (기존 매핑 함수에서 사용)
    
    Args:
        cell: python-docx Cell 객체
        content: 설정할 내용 (옵션)
    """
    if content is not None:
        # 기존 내용 제거
        cell.text = ""
        
        # 새 내용을 맑은 고딕으로 추가
        if '\n' in content:
            lines = content.split('\n')
            paragraph = cell.paragraphs[0]
            
            for i, line in enumerate(lines):
                if i > 0:
                    run = paragraph.add_run()
                    run.add_break()
                run = paragraph.add_run(line)
                apply_malgun_gothic_to_run(run)
        else:
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            paragraph.clear()
            run = paragraph.add_run(str(content))
            apply_malgun_gothic_to_run(run)
    else:
        # 기존 내용에 폰트만 적용
        for paragraph in cell.paragraphs:
            apply_malgun_gothic_to_paragraph(paragraph)


def apply_malgun_gothic_to_run(run) -> None:
    """
    Run 객체에 맑은 고딕 폰트 적용
    
    Args:
        run: python-docx Run 객체
    """
    if run.font:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)
        # 한글 폰트 설정
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')