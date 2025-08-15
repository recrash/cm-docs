"""
## A) Word 변경관리 문서 (template.docx)
- 템플릿: templates/template.docx (수정 금지, 로드만)
- python-docx로 **표 인덱스 좌표 주입**(VBA와 동일 의미/좌표)
  - Table(2)
    - Cell #9  ← writer_short (처리자_약칭)
    - Cell #17 ← created_date (없으면 오늘; 포맷 "mm/dd")
  - Table(3)  (1-base 인덱스, 행 순회 기준 셀 번호)
    - #2  ← title
    - #4  ← change_id
    - #6  ← work_datetime
    - #8  ← deploy_datetime
    - #10 ← customer
    - #12 ← request_dept
    - #14 ← requester
    - #18 ← doc_no
    - #20 ← system
    - #28 ← worker_deployer
    - #43 ← 목적-개선내용 (details.summary 또는 plan을 합성; 멀티라인은 run.add_break())
    - #46 ← impact_targets
    - #48 ← system
- 출력 파일명: `[YYMMDD 처리자] 변경관리요청서 {change_id} {title}.docx`
- 파일명 sanitize 규칙: 금지문자(\ / : * ? " < > |)만 `_`로 치환. 대괄호/괄호는 허용.
- ***인덱스 및 좌표 임의 변경 금지***
"""
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.text.run import Run

from ..models import ChangeRequest
from .paths import verify_template_exists, get_documents_dir
from .filename import generate_word_filename, unique_path
from .word_payload import build_word_payload


def _add_multiline_text(cell, text: str):
    """멀티라인 텍스트를 셀에 추가 (줄바꿈 처리)"""
    if not text:
        return
    
    # 기존 내용 제거
    for paragraph in cell.paragraphs:
        paragraph.clear()
    
    # 새 내용 추가
    paragraph = cell.paragraphs[0]
    lines = text.splitlines()
    
    for i, line in enumerate(lines):
        if i > 0:
            # 줄바꿈 추가
            run = paragraph.add_run()
            run.add_break()
        paragraph.add_run(line)


def _set_cell_text(cell, text: str):
    """셀에 텍스트 설정"""
    if text is None:
        text = ""
    
    # 멀티라인 텍스트인지 확인
    if '\n' in text:
        _add_multiline_text(cell, text)
    else:
        cell.text = str(text)


def build_change_request_doc(data: ChangeRequest, out_dir: Optional[Path] = None, raw_data: Optional[dict] = None) -> Path:
    """변경관리 요청서 Word 문서 생성
    
    Args:
        data: 변경 요청 데이터
        out_dir: 출력 디렉터리 (None이면 기본 documents 디렉터리 사용)
        raw_data: 원본 파싱 데이터 (필드 매핑용)
        
    Returns:
        Path: 생성된 파일 경로
        
    Raises:
        FileNotFoundError: 템플릿 파일이 없는 경우
        ValueError: 필수 데이터가 없는 경우
    """
    # 필수 데이터 검증
    if not data.change_id:
        raise ValueError("change_id는 필수입니다")
    if not data.title:
        raise ValueError("title은 필수입니다")
    
    # 템플릿 로드
    template_path = verify_template_exists("template.docx")
    doc = Document(str(template_path))
    
    # 출력 디렉터리 설정
    if out_dir is None:
        out_dir = get_documents_dir()
    
    # 원본 데이터가 있으면 Word용 payload 생성 (누락 필드 보완)
    if raw_data:
        # 원본 데이터를 Word 템플릿용으로 변환
        word_payload = build_word_payload(raw_data)
        
        # Word payload에서 필요한 값들 추출
        customer = word_payload.get("고객사", data.customer or "")
        worker_deployer = word_payload.get("작업자-배포자", data.worker_deployer or "")
        purpose_content = word_payload.get("목적-개선내용", "")
        impact_targets = word_payload.get("영향도_대상자", data.impact_targets or "")
        created_date = word_payload.get("작성일_mmdd", data.created_date or datetime.now().strftime("%m/%d"))
    else:
        # 기존 로직 (raw_data가 없는 경우)
        customer = data.customer or ""
        worker_deployer = data.worker_deployer or ""
        impact_targets = data.impact_targets or ""
        created_date = data.created_date or datetime.now().strftime("%m/%d")
        
        # 목적-개선내용 조합
        purpose_content = ""
        if data.details and data.details.summary:
            purpose_content = data.details.summary
        elif data.details and data.details.plan:
            purpose_content = data.details.plan
    
    try:
        # Table(2) 매핑
        table2 = doc.tables[1]  # 0-based index, Table(2) = index 1
        
        # Cell #9 (0-based index 8) ← writer_short
        table2._cells[8].text = data.writer_short or ""
        
        # Cell #17 (0-based index 16) ← created_date
        table2._cells[16].text = created_date
        
        # Table(3) 매핑
        table3 = doc.tables[2]  # 0-based index, Table(3) = index 2
        cells = table3._cells
        
        # 매핑 정보 (1-based → 0-based 변환) - VBA 매크로 요구사항에 따른 정확한 좌표
        mappings = {
            1: data.title,                    # #2 ← title
            3: data.change_id,                # #4 ← change_id
            5: data.work_datetime or "",      # #6 ← work_datetime
            7: data.deploy_datetime or "",    # #8 ← deploy_datetime
            9: customer,                      # #10 ← customer (파생 필드)
            11: data.request_dept or "",      # #12 ← request_dept
            13: data.requester or "",         # #14 ← requester
            17: data.doc_no or "",            # #18 ← doc_no
            19: data.system,                  # #20 ← system
            27: worker_deployer,              # #28 ← worker_deployer (파생 필드)
            42: purpose_content,              # #43 ← 목적-개선내용 (파생 필드)
            45: impact_targets,               # #46 ← impact_targets (파생 필드)
            47: data.system                   # #48 ← system (중복, VBA 매크로 요구사항)
        }
        
        # 매핑 적용
        for cell_index, value in mappings.items():
            if cell_index < len(cells):
                _set_cell_text(cells[cell_index], str(value))
        
    except IndexError as e:
        raise ValueError(f"템플릿 구조가 예상과 다릅니다: {e}")
    
    # 파일명 생성
    filename = generate_word_filename(
        change_id=data.change_id,
        title=data.title,
        writer_short=data.writer_short
    )
    
    # 중복 방지 경로 생성
    output_path = unique_path(out_dir, filename)
    
    # 문서 저장
    doc.save(str(output_path))
    
    return output_path