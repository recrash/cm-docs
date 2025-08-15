"""
Word 빌더 테스트

python-docx로 생성된 Word 문서의 표 좌표 매핑 검증
"""
import json
import pytest
import tempfile
from pathlib import Path
from datetime import datetime

from docx import Document

from ..models import ChangeRequest, ChangeDetails
from ..services.word_builder import build_change_request_doc


class TestWordBuilder:
    """Word 빌더 테스트"""
    
    @pytest.fixture
    def sample_change_request(self):
        """테스트용 변경 요청 데이터"""
        fixtures_dir = Path(__file__).parent / "fixtures"
        json_file = fixtures_dir / "itsupp_sample.json"
        
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # JSON 데이터를 ChangeRequest로 변환
        return ChangeRequest(
            change_id=data["변경관리번호"],
            system=data["시스템"],
            system_short=data["시스템_약칭"],
            title=data["제목"],
            requester=data["요청자"],
            request_dept=data["요청부서"],
            customer=data.get("고객사", ""),
            writer_short=data["처리자_약칭"],
            doc_no=data["문서번호"],
            work_datetime=data["작업일시"],
            deploy_datetime=data["배포일시"],
            biz_test_date=data.get("현업_테스트_일자", ""),
            details=ChangeDetails(
                summary=data.get("요구사항 상세분석", ""),
                plan=data.get("의뢰내용", "")
            ),
            impact_targets=data.get("영향도_대상자", ""),
            worker_deployer=data.get("작업자-배포자", ""),
            created_date=data.get("작성일", "")
        )
    
    @pytest.fixture
    def temp_output_dir(self):
        """임시 출력 디렉터리"""
        with tempfile.TemporaryDirectory() as temp_dir:
            yield Path(temp_dir)
    
    def test_build_word_document_success(self, sample_change_request, temp_output_dir):
        """Word 문서 생성 성공 테스트"""
        # Word 문서 생성
        output_path = build_change_request_doc(sample_change_request, temp_output_dir)
        
        # 파일이 생성되었는지 확인
        assert output_path.exists(), "Word 문서가 생성되지 않았습니다"
        assert output_path.suffix == '.docx', "파일 확장자가 .docx가 아닙니다"
        
        # 파일 크기 확인 (0바이트가 아님)
        assert output_path.stat().st_size > 0, "생성된 Word 문서가 비어있습니다"
    
    def test_word_document_filename_format(self, sample_change_request, temp_output_dir):
        """Word 문서 파일명 형식 검증"""
        output_path = build_change_request_doc(sample_change_request, temp_output_dir)
        
        filename = output_path.name
        
        # 파일명 형식: [YYMMDD 처리자] 변경관리요청서 {change_id} {title}.docx
        assert filename.startswith('['), "파일명이 대괄호로 시작하지 않습니다"
        assert '변경관리요청서' in filename, "파일명에 '변경관리요청서'가 없습니다"
        assert sample_change_request.change_id in filename, "파일명에 change_id가 없습니다"
        assert filename.endswith('.docx'), "파일명이 .docx로 끝나지 않습니다"
        
        # 날짜 형식 확인 (YYMMDD)
        date_part = filename[1:7]  # [YYMMDD 부분
        assert len(date_part) == 6, "날짜 부분 길이가 6자가 아닙니다"
        assert date_part.isdigit(), "날짜 부분이 숫자가 아닙니다"
    
    def test_table2_cell_mapping(self, sample_change_request, temp_output_dir):
        """Table(2) 셀 매핑 검증"""
        output_path = build_change_request_doc(sample_change_request, temp_output_dir)
        
        # 생성된 문서 열기
        doc = Document(str(output_path))
        
        # Table(2) = index 1
        assert len(doc.tables) >= 2, "문서에 필요한 테이블이 부족합니다"
        table2 = doc.tables[1]
        
        # Cell #9 (0-based index 8) ← writer_short
        cell_8 = table2._cells[8]
        assert cell_8.text.strip() == (sample_change_request.writer_short or ""), (
            f"Cell #9 값이 일치하지 않습니다. "
            f"실제: '{cell_8.text}', 기대: '{sample_change_request.writer_short or ''}'"
        )
        
        # Cell #17 (0-based index 16) ← created_date
        cell_16 = table2._cells[16]
        expected_date = sample_change_request.created_date or datetime.now().strftime("%m/%d")
        assert cell_16.text.strip() == expected_date, (
            f"Cell #17 값이 일치하지 않습니다. "
            f"실제: '{cell_16.text}', 기대: '{expected_date}'"
        )
    
    def test_table3_cell_mapping(self, sample_change_request, temp_output_dir):
        """Table(3) 셀 매핑 검증"""
        output_path = build_change_request_doc(sample_change_request, temp_output_dir)
        
        # 생성된 문서 열기
        doc = Document(str(output_path))
        
        # Table(3) = index 2
        assert len(doc.tables) >= 3, "문서에 필요한 테이블이 부족합니다"
        table3 = doc.tables[2]
        cells = table3._cells
        
        # 주요 매핑 검증 (1-based → 0-based 변환)
        test_mappings = [
            (1, sample_change_request.title, "title"),          # #2 ← title
            (3, sample_change_request.change_id, "change_id"),  # #4 ← change_id
            (13, sample_change_request.requester or "", "requester"),  # #14 ← requester
            (17, sample_change_request.doc_no or "", "doc_no"), # #18 ← doc_no
            (19, sample_change_request.system, "system")        # #20 ← system
        ]
        
        for cell_index, expected_value, field_name in test_mappings:
            if cell_index < len(cells):
                actual_value = cells[cell_index].text.strip()
                assert actual_value == str(expected_value), (
                    f"Cell #{cell_index + 1} ({field_name}) 값이 일치하지 않습니다. "
                    f"실제: '{actual_value}', 기대: '{expected_value}'"
                )
    
    def test_multiline_text_handling(self, temp_output_dir):
        """멀티라인 텍스트 처리 테스트"""
        # 멀티라인 텍스트가 포함된 테스트 데이터
        multiline_request = ChangeRequest(
            change_id="TEST_20250814_1",
            system="테스트시스템",
            title="멀티라인 테스트",
            details=ChangeDetails(
                summary="첫 번째 줄\n두 번째 줄\n세 번째 줄"
            )
        )
        
        output_path = build_change_request_doc(multiline_request, temp_output_dir)
        
        # 문서 열기 및 멀티라인 텍스트 확인
        doc = Document(str(output_path))
        table3 = doc.tables[2]
        
        # Cell #43 (0-based index 42) ← 목적-개선내용
        if len(table3._cells) > 42:
            purpose_cell = table3._cells[42]
            cell_text = purpose_cell.text
            
            # 줄바꿈이 포함되어 있는지 확인
            assert "첫 번째 줄" in cell_text, "첫 번째 줄이 포함되지 않았습니다"
            assert "두 번째 줄" in cell_text, "두 번째 줄이 포함되지 않았습니다"
            assert "세 번째 줄" in cell_text, "세 번째 줄이 포함되지 않았습니다"
    
    def test_missing_required_fields_error(self, temp_output_dir):
        """필수 필드 누락 시 에러 발생 테스트"""
        # change_id 누락
        with pytest.raises(ValueError, match="change_id는 필수입니다"):
            build_change_request_doc(
                ChangeRequest(change_id="", system="테스트", title="테스트"),
                temp_output_dir
            )
        
        # title 누락
        with pytest.raises(ValueError, match="title은 필수입니다"):
            build_change_request_doc(
                ChangeRequest(change_id="TEST_001", system="테스트", title=""),
                temp_output_dir
            )
    
    def test_missing_template_error(self, sample_change_request, temp_output_dir, monkeypatch):
        """템플릿 파일 누락 시 에러 발생 테스트"""
        # 존재하지 않는 템플릿 경로로 변경
        def mock_verify_template_exists(template_name):
            raise FileNotFoundError(f"템플릿 파일을 찾을 수 없습니다: {template_name}")
        
        monkeypatch.setattr(
            "autodoc_service.app.services.word_builder.verify_template_exists", 
            mock_verify_template_exists
        )
        
        with pytest.raises(FileNotFoundError, match="템플릿 파일을 찾을 수 없습니다"):
            build_change_request_doc(sample_change_request, temp_output_dir)
    
    def test_created_date_default_value(self, temp_output_dir):
        """created_date 기본값 테스트"""
        request_without_date = ChangeRequest(
            change_id="TEST_001",
            system="테스트시스템",
            title="날짜 테스트"
        )
        
        output_path = build_change_request_doc(request_without_date, temp_output_dir)
        
        # 문서 열기
        doc = Document(str(output_path))
        table2 = doc.tables[1]
        
        # Cell #17에 오늘 날짜 mm/dd 형식이 들어있는지 확인
        cell_16 = table2._cells[16]
        today_mmdd = datetime.now().strftime("%m/%d")
        
        assert cell_16.text.strip() == today_mmdd, (
            f"기본 작성일이 올바르지 않습니다. "
            f"실제: '{cell_16.text}', 기대: '{today_mmdd}'"
        )
    
    def test_file_path_different_from_template(self, sample_change_request, temp_output_dir):
        """생성된 파일이 템플릿과 다른 경로인지 확인"""
        output_path = build_change_request_doc(sample_change_request, temp_output_dir)
        
        # 출력 파일이 documents/ 디렉터리에 있는지 확인
        assert output_path.parent == temp_output_dir, (
            f"출력 파일이 지정된 디렉터리에 생성되지 않았습니다. "
            f"실제: {output_path.parent}, 기대: {temp_output_dir}"
        )
        
        # 파일명이 템플릿 파일명과 다른지 확인
        assert output_path.name != "template.docx", (
            "생성된 파일명이 템플릿 파일명과 동일합니다. 새로운 파일명이어야 합니다."
        )