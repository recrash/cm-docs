"""
Word 라벨 기반 빌더 테스트

라벨 기반 매핑으로 생성된 문서에 핵심 값들이 포함되는지 검증합니다.
"""
from pathlib import Path
import json
import tempfile

from docx import Document

from ..models import ChangeRequest
from ..services.label_based_word_builder import build_change_request_doc_label_based


def _collect_all_texts(doc: Document) -> list[str]:
    texts: list[str] = []
    # 본문 단락
    for p in doc.paragraphs:
        if p.text:
            texts.append(p.text)
    # 테이블 셀
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text:
                    texts.append(c.text)
    # 헤더/푸터
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                if p.text:
                    texts.append(p.text)
        if section.footer:
            for p in section.footer.paragraphs:
                if p.text:
                    texts.append(p.text)
    return texts


class TestWordLabelBuilder:
    def test_build_word_doc_and_mapping_contains_core_values(self):
        # 1) 픽스처 로드
        fixtures_dir = Path(__file__).parent / "fixtures"
        json_file = fixtures_dir / "itsupp_sample.json"
        with open(json_file, "r", encoding="utf-8") as f:
            raw = json.load(f)

        # ChangeRequest(필수 최소값)
        change_id = raw.get("변경관리번호") or "LIMS_20250728_1"
        title = raw.get("제목") or "테스트 제목"
        system = raw.get("시스템") or "LIMS-001"
        requester = raw.get("요청자") or "홍길동"

        change_req = ChangeRequest(
            change_id=change_id,
            system=system,
            title=title,
            requester=requester,
            writer_short=raw.get("처리자_약칭", None),
        )

        # 2) 문서 생성
        with tempfile.TemporaryDirectory() as temp_dir:
            out_dir = Path(temp_dir)
            output_path = build_change_request_doc_label_based(change_req, out_dir=out_dir, raw_data=raw)

            # 3) 파일 생성 확인
            assert output_path.exists(), "Word 문서가 생성되지 않았습니다"
            assert output_path.suffix == ".docx", "확장자가 .docx가 아닙니다"
            assert output_path.stat().st_size > 0, "생성된 Word 문서가 비어있습니다"

            # 4) 문서 열어 내용 확인
            doc = Document(str(output_path))
            texts = "\n".join(_collect_all_texts(doc))

            # 핵심 값들이 문서 어딘가에 존재해야 함 (라벨 기반 매핑 결과)
            assert title in texts, "제목이 문서에 존재하지 않습니다"
            assert change_id in texts, "변경관리번호가 문서에 존재하지 않습니다"

            # 목적/개선내용(구조화 텍스트) 일부 키워드 검사
            assert "1. 목적" in texts, "구조화된 목적 섹션이 없습니다"
            assert "2. 주요내용" in texts, "구조화된 주요내용 섹션이 없습니다"

            # 요청자/요청부서 또는 고객사 중 일부 값 존재 확인(템플릿에 따라 어느 셀에 들어가도 텍스트 포함되어야 함)
            req_dept = raw.get("요청부서", "")
            customer = raw.get("고객사", "")
            assert requester in texts, "요청자가 문서에 존재하지 않습니다"
            assert (req_dept in texts) or (customer in texts) or (raw.get("신청자", "") in texts), (
                "요청부서/고객사/신청자 중 어느 것도 문서에 존재하지 않습니다"
            )

            # 작업/배포 일시(파서 파생 규칙)
            work_dt = raw.get("작업일시", "")
            deploy_dt = raw.get("배포일시", "")
            if work_dt:
                assert work_dt in texts, "작업일시가 문서에 존재하지 않습니다"
            if deploy_dt:
                assert deploy_dt in texts, "배포일시가 문서에 존재하지 않습니다"


