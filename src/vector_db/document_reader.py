import os
import pandas as pd
from docx import Document
from typing import Dict, List, Any, Optional
import logging

class DocumentReader:
    """다양한 문서 형식을 읽는 클래스"""
    
    def __init__(self):
        self.supported_extensions = {'.xlsx', '.xls', '.docx', '.txt', '.md'}
        
    def read_document(self, file_path: str) -> Dict[str, Any]:
        """
        파일을 읽어서 텍스트 내용을 반환
        
        Args:
            file_path: 읽을 파일 경로
            
        Returns:
            문서 정보 딕셔너리 (content, metadata)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext in ['.xlsx', '.xls']:
                return self._read_excel(file_path)
            elif file_ext == '.docx':
                return self._read_docx(file_path)
            elif file_ext in ['.txt', '.md']:
                return self._read_text(file_path)
            else:
                raise ValueError(f"지원하지 않는 파일 형식입니다: {file_ext}")
                
        except Exception as e:
            logging.error(f"파일 읽기 중 오류 발생 ({file_path}): {e}")
            return {
                'content': f"파일 읽기 오류: {str(e)}",
                'metadata': {
                    'file_path': file_path,
                    'file_type': file_ext,
                    'error': str(e),
                    'status': 'error'
                }
            }
    
    def _read_excel(self, file_path: str) -> Dict[str, Any]:
        """엑셀 파일 읽기"""
        try:
            # 모든 시트 읽기
            excel_file = pd.ExcelFile(file_path)
            all_content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # 시트 헤더 추가
                all_content.append(f"=== 시트: {sheet_name} ===")
                
                # 컬럼명 추가
                if not df.empty:
                    all_content.append(f"컬럼: {', '.join(df.columns.astype(str))}")
                    
                    # 데이터 내용을 텍스트로 변환
                    for index, row in df.iterrows():
                        row_text = []
                        for col_name, value in row.items():
                            if pd.notna(value) and str(value).strip():
                                row_text.append(f"{col_name}: {str(value).strip()}")
                        
                        if row_text:
                            all_content.append(" | ".join(row_text))
                else:
                    all_content.append("(빈 시트)")
                
                all_content.append("")  # 시트 간 구분
            
            content = "\n".join(all_content)
            
            return {
                'content': content,
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'excel',
                    'sheet_count': len(excel_file.sheet_names),
                    'sheet_names': excel_file.sheet_names,
                    'total_rows': sum(pd.read_excel(file_path, sheet_name=sheet).shape[0] 
                                    for sheet in excel_file.sheet_names),
                    'status': 'success'
                }
            }
            
        except Exception as e:
            raise Exception(f"엑셀 파일 읽기 실패: {e}")
    
    def _read_docx(self, file_path: str) -> Dict[str, Any]:
        """워드 파일 읽기"""
        try:
            doc = Document(file_path)
            all_content = []
            
            # 문단 읽기
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    all_content.append(text)
            
            # 표 읽기
            if doc.tables:
                all_content.append("\n=== 표 데이터 ===")
                for table_idx, table in enumerate(doc.tables):
                    all_content.append(f"\n표 {table_idx + 1}:")
                    
                    for row_idx, row in enumerate(table.rows):
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_data.append(cell_text)
                        
                        if row_data:
                            all_content.append(" | ".join(row_data))
            
            content = "\n".join(all_content)
            
            return {
                'content': content,
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'docx',
                    'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                    'table_count': len(doc.tables),
                    'status': 'success'
                }
            }
            
        except Exception as e:
            raise Exception(f"워드 파일 읽기 실패: {e}")
    
    def _read_text(self, file_path: str) -> Dict[str, Any]:
        """텍스트 파일 읽기"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                'content': content,
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'text',
                    'char_count': len(content),
                    'line_count': len(content.splitlines()),
                    'status': 'success'
                }
            }
            
        except UnicodeDecodeError:
            # UTF-8로 읽기 실패시 다른 인코딩 시도
            try:
                with open(file_path, 'r', encoding='cp949') as f:
                    content = f.read()
                return {
                    'content': content,
                    'metadata': {
                        'file_path': file_path,
                        'file_type': 'text',
                        'encoding': 'cp949',
                        'char_count': len(content),
                        'line_count': len(content.splitlines()),
                        'status': 'success'
                    }
                }
            except Exception as e:
                raise Exception(f"텍스트 파일 인코딩 오류: {e}")
        
        except Exception as e:
            raise Exception(f"텍스트 파일 읽기 실패: {e}")
    
    def is_supported_file(self, file_path: str) -> bool:
        """지원하는 파일 형식인지 확인"""
        file_ext = os.path.splitext(file_path)[1].lower()
        return file_ext in self.supported_extensions
    
    def get_supported_extensions(self) -> List[str]:
        """지원하는 파일 확장자 목록 반환"""
        return list(self.supported_extensions)